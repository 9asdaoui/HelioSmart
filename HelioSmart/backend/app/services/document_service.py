"""Document Processing Service - handles file uploads and LLM-based extraction"""
import os
import json
import uuid
import shutil
from typing import Optional, List, Dict, Any, Tuple
from pathlib import Path
from datetime import datetime
import httpx

from sqlalchemy.orm import Session
from fastapi import UploadFile, HTTPException, status

from app.core.config import settings
from app.models.user import Vendor, VendorDocument, DocumentStatus, Product, ProductStatus
from app.schemas.auth import DocumentExtractionResponse, ProductCreate


# Configuration
UPLOAD_DIR = Path("uploads/vendor_documents")
ALLOWED_EXTENSIONS = {'.pdf', '.csv', '.xlsx', '.xls', '.doc', '.docx', '.txt'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

# Ensure upload directory exists
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


class DocumentService:
    """Service for handling vendor document uploads and processing"""
    
    def __init__(self, db: Session):
        self.db = db
        self.llm_service = LLMExtractionService()
    
    def validate_file(self, file: UploadFile) -> Tuple[bool, str]:
        """Validate uploaded file"""
        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            return False, f"Invalid file type. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
        
        # Check content type
        allowed_types = [
            'application/pdf',
            'text/csv',
            'application/vnd.ms-excel',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain'
        ]
        
        if file.content_type not in allowed_types:
            return False, f"Invalid content type: {file.content_type}"
        
        return True, ""
    
    def save_upload_file(self, file: UploadFile, vendor_id: int) -> Tuple[str, str, int]:
        """Save uploaded file to disk and return path, filename, size"""
        # Generate unique filename
        file_ext = Path(file.filename).suffix.lower()
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        
        # Create vendor subdirectory
        vendor_dir = UPLOAD_DIR / str(vendor_id)
        vendor_dir.mkdir(exist_ok=True)
        
        file_path = vendor_dir / unique_filename
        
        # Save file
        try:
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            file_size = file_path.stat().st_size
            
            if file_size > MAX_FILE_SIZE:
                file_path.unlink()  # Delete file
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File too large. Maximum size: {MAX_FILE_SIZE / (1024*1024)}MB"
                )
            
            return str(file_path), unique_filename, file_size
            
        except Exception as e:
            if file_path.exists():
                file_path.unlink()
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to save file: {str(e)}"
            )
    
    def upload_document(
        self,
        vendor: Vendor,
        file: UploadFile,
        uploaded_by_id: Optional[int] = None
    ) -> VendorDocument:
        """Upload and save a vendor document"""
        # Validate file
        is_valid, error_msg = self.validate_file(file)
        if not is_valid:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=error_msg
            )
        
        # Save file
        file_path, unique_filename, file_size = self.save_upload_file(file, vendor.id)
        
        # Create document record
        document = VendorDocument(
            vendor_id=vendor.id,
            uploaded_by_id=uploaded_by_id,
            filename=unique_filename,
            original_filename=file.filename,
            file_path=file_path,
            file_size=file_size,
            file_type=file.content_type,
            file_extension=Path(file.filename).suffix.lower(),
            status=DocumentStatus.UPLOADED
        )
        
        self.db.add(document)
        
        # Update vendor document count
        vendor.total_documents += 1
        
        self.db.commit()
        self.db.refresh(document)
        
        return document
    
    def process_document(self, document: VendorDocument) -> DocumentExtractionResponse:
        """Process document with LLM to extract product data"""
        # Update status to processing
        document.status = DocumentStatus.PROCESSING
        self.db.commit()
        
        try:
            # Extract text from document
            extracted_text = self._extract_text_from_document(document)
            document.extracted_text = extracted_text
            
            # Use LLM to extract structured data
            extraction_result = self.llm_service.extract_products(
                extracted_text,
                document.file_extension
            )
            
            # Update document with extraction results
            document.extracted_data = extraction_result.get("raw_data")
            document.extraction_confidence = extraction_result.get("confidence", 0.0)
            document.status = DocumentStatus.EXTRACTED
            document.processed_at = datetime.utcnow()
            
            # Create product records from extraction
            products = self._create_products_from_extraction(
                document.vendor_id,
                document.id,
                extraction_result.get("products", [])
            )
            
            self.db.commit()
            
            return DocumentExtractionResponse(
                document_id=document.id,
                status=document.status,
                extracted_data=document.extracted_data,
                extraction_confidence=document.extraction_confidence,
                extracted_product_count=len(products),
                processing_error=None
            )
            
        except Exception as e:
            document.status = DocumentStatus.ERROR
            document.processing_error = str(e)
            self.db.commit()
            
            return DocumentExtractionResponse(
                document_id=document.id,
                status=DocumentStatus.ERROR,
                extracted_data=None,
                extraction_confidence=None,
                extracted_product_count=0,
                processing_error=str(e)
            )
    
    def _extract_text_from_document(self, document: VendorDocument) -> str:
        """Extract text content from various document types"""
        file_path = Path(document.file_path)
        
        if document.file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif document.file_extension == '.pdf':
            try:
                import PyPDF2
                text = ""
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                return text
            except ImportError:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="PDF processing not available. Install PyPDF2."
                )
        
        elif document.file_extension in ['.csv']:
            import pandas as pd
            df = pd.read_csv(file_path)
            return df.to_string()
        
        elif document.file_extension in ['.xlsx', '.xls']:
            import pandas as pd
            df = pd.read_excel(file_path)
            return df.to_string()
        
        elif document.file_extension in ['.doc', '.docx']:
            try:
                import docx
                doc = docx.Document(file_path)
                return "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Word document processing not available. Install python-docx."
                )
        
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {document.file_extension}"
            )
    
    def _create_products_from_extraction(
        self,
        vendor_id: int,
        document_id: int,
        extracted_products: List[Dict[str, Any]]
    ) -> List[Product]:
        """Create product records from LLM extraction"""
        products = []
        
        for product_data in extracted_products:
            try:
                product = Product(
                    vendor_id=vendor_id,
                    source_document_id=document_id,
                    name=product_data.get("name", "Unknown Product"),
                    sku=product_data.get("sku"),
                    description=product_data.get("description"),
                    category=product_data.get("category", "uncategorized"),
                    subcategory=product_data.get("subcategory"),
                    price=str(product_data.get("price")) if product_data.get("price") else None,
                    currency=product_data.get("currency", "MAD"),
                    unit=product_data.get("unit"),
                    specifications=product_data.get("specifications", {}),
                    status=ProductStatus.PENDING,
                    is_active=True,
                    extracted_data=product_data
                )
                
                self.db.add(product)
                products.append(product)
                
                # Update vendor product count
                vendor = self.db.query(Vendor).filter(Vendor.id == vendor_id).first()
                if vendor:
                    vendor.total_products += 1
                
            except Exception as e:
                # Log error but continue with other products
                print(f"Failed to create product from extraction: {e}")
                continue
        
        self.db.commit()
        return products
    
    def approve_extracted_products(
        self,
        document: VendorDocument,
        approved_product_ids: List[int],
        rejected_product_ids: List[int],
        notes: Optional[str] = None
    ) -> List[Product]:
        """Vendor approves or rejects extracted products"""
        approved_products = []
        
        # Update approved products
        for product_id in approved_product_ids:
            product = self.db.query(Product).filter(
                Product.id == product_id,
                Product.source_document_id == document.id
            ).first()
            
            if product:
                product.status = ProductStatus.APPROVED
                product.approved_at = datetime.utcnow()
                approved_products.append(product)
        
        # Update rejected products
        for product_id in rejected_product_ids:
            product = self.db.query(Product).filter(
                Product.id == product_id,
                Product.source_document_id == document.id
            ).first()
            
            if product:
                product.status = ProductStatus.REJECTED
        
        # Update document
        document.is_vendor_approved = True
        document.vendor_approved_at = datetime.utcnow()
        document.vendor_notes = notes
        document.status = DocumentStatus.APPROVED
        
        self.db.commit()
        
        return approved_products
    
    def get_vendor_documents(self, vendor_id: int) -> List[VendorDocument]:
        """Get all documents for a vendor"""
        return self.db.query(VendorDocument).filter(
            VendorDocument.vendor_id == vendor_id
        ).order_by(VendorDocument.uploaded_at.desc()).all()
    
    def get_document_by_id(self, document_id: int, vendor_id: Optional[int] = None) -> Optional[VendorDocument]:
        """Get document by ID, optionally filtered by vendor"""
        query = self.db.query(VendorDocument).filter(VendorDocument.id == document_id)
        
        if vendor_id:
            query = query.filter(VendorDocument.vendor_id == vendor_id)
        
        return query.first()


class LLMExtractionService:
    """Service for LLM-based product extraction from documents"""
    
    def __init__(self):
        # For this implementation, we'll use a mock/simulation
        # In production, integrate with OpenAI, Claude, or other LLM API
        self.api_key = os.getenv("LLM_API_KEY", "")
        self.api_url = os.getenv("LLM_API_URL", "https://api.openai.com/v1/chat/completions")
        self.model = os.getenv("LLM_MODEL", "gpt-4")
    
    def extract_products(self, text: str, file_type: str) -> Dict[str, Any]:
        """Extract product data from document text using LLM"""
        
        # If LLM API is not configured, use mock extraction
        if not self.api_key:
            return self._mock_extraction(text, file_type)
        
        # Prepare prompt
        prompt = self._build_extraction_prompt(text, file_type)
        
        try:
            # Call LLM API
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": self.model,
                "messages": [
                    {
                        "role": "system",
                        "content": "You are a product data extraction specialist. Extract product information from documents and return structured JSON data."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "temperature": 0.1,
                "response_format": {"type": "json_object"}
            }
            
            with httpx.Client() as client:
                response = client.post(
                    self.api_url,
                    headers=headers,
                    json=payload,
                    timeout=60.0
                )
                response.raise_for_status()
                
                result = response.json()
                content = result["choices"][0]["message"]["content"]
                extracted_data = json.loads(content)
                
                return {
                    "products": extracted_data.get("products", []),
                    "confidence": extracted_data.get("confidence", 0.8),
                    "raw_data": extracted_data
                }
                
        except Exception as e:
            print(f"LLM extraction failed: {e}")
            # Fallback to mock extraction
            return self._mock_extraction(text, file_type)
    
    def _build_extraction_prompt(self, text: str, file_type: str) -> str:
        """Build the extraction prompt for the LLM"""
        return f"""
Extract product information from the following document content and return it as JSON.

Document Type: {file_type}

Document Content:
{text[:15000]}  # Limit text length

Extract the following fields for each product:
- name: Product name (required)
- sku: Product SKU or model number
- description: Product description
- category: Product category (e.g., panels, inverters, batteries, mounting, accessories)
- subcategory: More specific category
- price: Price as number or string
- currency: Currency code (default: MAD)
- unit: Unit of measurement (e.g., per unit, per watt, per set)
- specifications: Object with technical specs (power, voltage, dimensions, etc.)

Return JSON in this format:
{{
    "products": [
        {{
            "name": "...",
            "sku": "...",
            "description": "...",
            "category": "...",
            "subcategory": "...",
            "price": "...",
            "currency": "MAD",
            "unit": "...",
            "specifications": {{...}}
        }}
    ],
    "confidence": 0.95
}}
"""
    
    def _mock_extraction(self, text: str, file_type: str) -> Dict[str, Any]:
        """Mock extraction for testing without LLM API"""
        # Simulate product extraction with sample data
        # In real implementation, this would use the actual LLM
        
        products = []
        
        # Simple pattern matching for demo purposes
        lines = text.split('\n')
        
        for i, line in enumerate(lines[:20]):  # Check first 20 lines
            line_lower = line.lower()
            
            # Detect panels
            if any(keyword in line_lower for keyword in ['panel', 'module', 'watt', 'wp']):
                products.append({
                    "name": f"Solar Panel {i+1}",
                    "sku": f"PANEL-{i+1:03d}",
                    "description": line.strip()[:200],
                    "category": "panels",
                    "subcategory": "monocrystalline",
                    "price": None,
                    "currency": "MAD",
                    "unit": "per watt",
                    "specifications": {
                        "power": "400W",
                        "efficiency": "21%",
                        "voltage": "40V"
                    }
                })
            
            # Detect inverters
            elif any(keyword in line_lower for keyword in ['inverter', 'convertisseur', 'onduleur']):
                products.append({
                    "name": f"Solar Inverter {i+1}",
                    "sku": f"INV-{i+1:03d}",
                    "description": line.strip()[:200],
                    "category": "inverters",
                    "subcategory": "string inverter",
                    "price": None,
                    "currency": "MAD",
                    "unit": "per unit",
                    "specifications": {
                        "power": "5kW",
                        "input_voltage": "600V",
                        "efficiency": "98%"
                    }
                })
        
        return {
            "products": products if products else [
                {
                    "name": "Sample Solar Product",
                    "sku": "SAMPLE-001",
                    "description": "Automatically extracted sample product",
                    "category": "panels",
                    "subcategory": "monocrystalline",
                    "price": None,
                    "currency": "MAD",
                    "unit": "per watt",
                    "specifications": {}
                }
            ],
            "confidence": 0.75 if products else 0.5,
            "raw_data": {"note": "Mock extraction - configure LLM_API_KEY for real extraction"}
        }
